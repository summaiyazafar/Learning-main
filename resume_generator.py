"""
Professional Resume Generator
AI Resume Tailoring System

Purpose:
- Generate ATS-friendly DOCX resume
- Preserve protected personal information
- Render only supplied information
- Work directly with ResumeTailor output
- Never add missing JD skills to candidate skills
- Never fabricate employment, education, certifications,
  projects, or experience
"""

import os
import re
from copy import deepcopy
from typing import Any, Dict, List, Optional, Set, Union

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class ResumeGenerator:
    """
    Professional resume generator producing DOCX output.

    All protected personal information is preserved exactly as supplied.
    Missing JD skills are never added to the final skills section.
    """

    # ============================================================
    # INITIALIZATION
    # ============================================================

    def __init__(self) -> None:
        """Initialize the generator with default font and protected fields."""
        self.default_font: str = "Arial"

        # These fields must never be changed by tailoring.
        self.protected_fields: List[str] = [
            "name",
            "phone",
            "email",
            "linkedin",
            "github",
            "kaggle",
            "location",
            "education",
            "certifications",
        ]

    # ============================================================
    # TEXT CLEANING
    # ============================================================

    @staticmethod
    def clean_text(text: Any) -> str:
        """
        Clean and normalize text for display.

        Protected values are stored separately and are never
        semantically rewritten. This only removes extra whitespace.
        """
        if text is None:
            return ""
        text = str(text)
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    # ============================================================
    # SAFE VALUE
    # ============================================================

    @staticmethod
    def safe_value(value: Any, default: str = "") -> str:
        """Return a safe string representation of the value."""
        if value is None:
            return default
        return str(value).strip()

    # ============================================================
    # NORMALIZE LIST
    # ============================================================

    def normalize_list(self, value: Any) -> List[Union[str, Dict]]:
        """
        Convert different input types into a clean list.

        Supports:
            - None → []
            - str → split by newline, clean each line
            - list/tuple/set → clean each item, preserve dicts
            - any other → single-item list if non-empty
        """
        if value is None:
            return []

        if isinstance(value, str):
            result = []
            for line in value.splitlines():
                line = self.clean_text(line)
                if line:
                    result.append(line)
            return result

        if isinstance(value, (list, tuple, set)):
            result = []
            for item in value:
                if isinstance(item, dict):
                    result.append(deepcopy(item))
                    continue
                item = self.clean_text(item)
                if item:
                    result.append(item)
            return result

        return [value]

    # ============================================================
    # BULLET CONVERSION
    # ============================================================

    def convert_to_bullets(self, text: Any) -> List[str]:
        """
        Convert supplied content into clean bullet strings.

        Does not generate new information; only formats existing data.
        """
        if not text:
            return []

        if isinstance(text, (list, tuple, set)):
            lines = []
            for item in text:
                if isinstance(item, dict):
                    continue
                lines.append(str(item))
            text = "\n".join(lines)

        text = str(text)
        bullets = []
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            # Remove existing bullet symbols
            line = re.sub(r"^[\s•▪◦●○*-]+\s*", "", line)
            line = self.clean_text(line)
            if line:
                bullets.append(line)
        return bullets

    # ============================================================
    # SKILL NORMALIZATION
    # ============================================================

    def normalize_skills(self, skills: Any) -> Dict[str, List[str]]:
        """
        Normalize supplied skills into categorized dictionary.

        Supports:
            - List: ["Python", "SQL"] → auto-categorized
            - String: "Python, SQL" → parsed and categorized
            - Dict: {"Programming": ["Python"], "Data": ["SQL"]}

        IMPORTANT: No new skills are created.
        """
        if not skills:
            return {}

        # Already categorized
        if isinstance(skills, dict):
            result = {}
            for category, skill_list in skills.items():
                category = self.clean_text(category)
                if not category:
                    continue
                if isinstance(skill_list, str):
                    values = [self.clean_text(s) for s in skill_list.split(",") if self.clean_text(s)]
                    skill_list = values
                elif isinstance(skill_list, (list, tuple, set)):
                    values = []
                    for item in skill_list:
                        if isinstance(item, dict):
                            continue
                        item = self.clean_text(item)
                        if item:
                            values.append(item)
                    skill_list = values
                else:
                    skill_list = []
                if skill_list:
                    result[category] = self.remove_duplicate_skills(skill_list)
            return result

        # String or list
        if isinstance(skills, str):
            skills = [self.clean_text(s) for s in skills.split(",") if self.clean_text(s)]
        elif isinstance(skills, (list, tuple, set)):
            skills = [self.clean_text(s) for s in skills if self.clean_text(s) and not isinstance(s, dict)]
        else:
            return {}

        return self.organize_skills(skills)

    # ============================================================
    # REMOVE DUPLICATES
    # ============================================================

    @staticmethod
    def remove_duplicate_skills(skills: List[str]) -> List[str]:
        """Remove duplicate skills while preserving order."""
        result = []
        seen = set()
        for skill in skills:
            key = str(skill).strip().lower()
            if not key:
                continue
            if key not in seen:
                seen.add(key)
                result.append(skill)
        return result

    # ============================================================
    # SKILL GROUPING
    # ============================================================

    def organize_skills(self, skills: List[str]) -> Dict[str, List[str]]:
        """
        Organize skills into categories based on predefined sets.

        IMPORTANT: This function NEVER adds skills; it only groups existing ones.
        """
        # Define category sets
        programming_set = {
            "python", "r", "java", "javascript", "typescript",
            "c", "c++", "c#", "go", "rust", "php", "ruby",
            "kotlin", "swift", "scala", "bash", "powershell",
        }

        data_set = {
            "sql", "excel", "power query", "power pivot", "vba",
            "pandas", "numpy", "scipy", "matplotlib", "seaborn",
            "plotly", "statsmodels", "jupyter", "anaconda",
            "data analysis", "data analytics", "data visualization",
            "exploratory data analysis", "eda", "data cleaning",
            "data preprocessing", "statistics", "statistical analysis",
            "hypothesis testing", "a/b testing", "regression analysis",
            "time series", "forecasting", "business intelligence",
            "dashboard", "dashboards", "reporting", "kpi", "kpis",
            "power bi", "powerbi", "tableau", "looker", "looker studio", "qlik",
        }

        ai_ml_set = {
            "artificial intelligence", "ai", "machine learning", "ml",
            "deep learning", "neural networks", "supervised learning",
            "unsupervised learning", "reinforcement learning",
            "scikit-learn", "sklearn", "tensorflow", "keras",
            "pytorch", "xgboost", "lightgbm", "catboost",
            "random forest", "decision tree", "logistic regression",
            "linear regression", "support vector machine", "svm",
            "knn", "k-nearest neighbors", "clustering", "k-means",
            "pca", "feature engineering", "model evaluation",
            "cross validation", "hyperparameter tuning",
            "ensemble learning", "anomaly detection",
            "recommendation systems", "nlp", "natural language processing",
            "text classification", "named entity recognition", "ner",
            "sentiment analysis", "tokenization", "embeddings",
            "transformers", "bert", "gpt", "llm", "llms",
            "large language models", "generative ai", "genai",
            "prompt engineering", "langchain", "llamaindex", "rag",
            "retrieval augmented generation", "vector database",
            "faiss", "pinecone", "chromadb", "openai",
            "hugging face", "huggingface", "fine tuning", "fine-tuning",
            "lora", "peft", "ai agents", "agents", "function calling",
            "computer vision", "opencv", "yolo", "object detection",
            "image classification", "image segmentation", "ocr",
            "cnn", "convolutional neural network", "image processing",
        }

        databases_set = {
            "mysql", "postgresql", "postgres", "sql server",
            "microsoft sql server", "oracle", "sqlite",
            "mongodb", "redis", "cassandra", "snowflake",
            "bigquery", "redshift", "databricks", "faiss",
            "pinecone", "chromadb",
        }

        web_set = {
            "html", "css", "javascript", "react", "node.js",
            "nodejs", "django", "flask", "fastapi", "streamlit",
            "rest api", "rest apis", "api", "apis",
        }

        tools_set = {
            "git", "github", "gitlab", "bitbucket", "docker",
            "kubernetes", "jupyter", "anaconda", "vs code",
            "visual studio code", "postman", "mlflow", "dvc",
        }

        cloud_set = {
            "aws", "amazon web services", "azure", "microsoft azure",
            "gcp", "google cloud", "aws s3", "ec2", "lambda",
            "azure blob storage", "azure machine learning",
            "azure ml", "sagemaker", "vertex ai",
        }

        data_engineering_set = {
            "apache spark", "spark", "apache kafka", "kafka",
            "airflow", "apache airflow", "dbt", "hadoop",
            "hive", "presto", "trino", "data pipelines",
            "etl", "elt", "data warehouse", "data lake",
            "data lakehouse", "delta lake",
        }

        soft_skills_set = {
            "communication", "written communication",
            "verbal communication", "leadership", "teamwork",
            "team collaboration", "problem solving",
            "problem-solving", "critical thinking",
            "time management", "adaptability", "creativity",
            "attention to detail", "presentation",
            "stakeholder management", "project management",
            "agile", "scrum",
        }

        business_set = {
            "business analysis", "business analytics",
            "requirements analysis", "market research",
            "sales", "marketing", "digital marketing",
            "seo", "sem", "content marketing", "crm",
            "customer relationship management",
            "financial analysis", "accounting",
            "budgeting", "risk management",
        }

        # Initialize buckets
        buckets = {
            "Programming": [],
            "Data & Analytics": [],
            "AI & Machine Learning": [],
            "Data Engineering": [],
            "Databases": [],
            "Web & APIs": [],
            "Tools & Development": [],
            "Cloud": [],
            "Business & Professional": [],
            "Soft Skills": [],
            "Other": [],
        }

        # Map skill to bucket
        for skill in skills:
            if isinstance(skill, dict):
                continue
            skill = self.clean_text(skill)
            if not skill:
                continue
            lower = skill.lower()
            if lower in programming_set:
                buckets["Programming"].append(skill)
            elif lower in data_set:
                buckets["Data & Analytics"].append(skill)
            elif lower in ai_ml_set:
                buckets["AI & Machine Learning"].append(skill)
            elif lower in data_engineering_set:
                buckets["Data Engineering"].append(skill)
            elif lower in databases_set:
                buckets["Databases"].append(skill)
            elif lower in web_set:
                buckets["Web & APIs"].append(skill)
            elif lower in tools_set:
                buckets["Tools & Development"].append(skill)
            elif lower in cloud_set:
                buckets["Cloud"].append(skill)
            elif lower in business_set:
                buckets["Business & Professional"].append(skill)
            elif lower in soft_skills_set:
                buckets["Soft Skills"].append(skill)
            else:
                buckets["Other"].append(skill)

        # Remove empty buckets and deduplicate
        result = {}
        for category, values in buckets.items():
            values = self.remove_duplicate_skills(values)
            if values:
                result[category] = values
        return result

    # ============================================================
    # SAFE SUMMARY
    # ============================================================

    def generate_summary(
        self,
        job_title: str,
        skills: Union[Dict[str, List[str]], List[str], str],
        experience: Optional[int] = None
    ) -> str:
        """
        Generate a conservative summary only from supplied data.

        This method intentionally avoids unsupported claims such as
        "2 years of experience", "expert", "highly experienced", etc.
        """
        job_title = self.clean_text(job_title) or "Professional"

        all_skills = []
        if isinstance(skills, dict):
            for skill_list in skills.values():
                all_skills.extend(self.normalize_list(skill_list))
        else:
            all_skills = self.normalize_list(skills)

        all_skills = self.remove_duplicate_skills(all_skills)
        selected = all_skills[:6]

        if selected:
            skill_text = ", ".join(selected)
            return (
                f"Professional targeting a {job_title} role "
                f"with skills in {skill_text}. "
                f"Focused on applying existing technical and "
                f"analytical capabilities to relevant projects "
                f"and practical problem solving."
            )
        return (
            f"Professional targeting a {job_title} role, "
            f"focused on applying existing skills to relevant "
            f"projects and practical problem solving."
        )

    # ============================================================
    # FORMATTING HELPERS
    # ============================================================

    def format_experience(self, experience: Any) -> List[Union[Dict, str]]:
        """Format experience entries into a consistent structure."""
        if not experience:
            return []
        if isinstance(experience, str):
            return self.convert_to_bullets(experience)
        if isinstance(experience, (list, tuple, set)):
            formatted = []
            for item in experience:
                if isinstance(item, dict):
                    entry = deepcopy(item)
                    bullets = entry.get("bullets") or entry.get("description") or entry.get("responsibilities") or []
                    entry["bullets"] = self.convert_to_bullets(bullets)
                    formatted.append(entry)
                else:
                    item = self.clean_text(item)
                    if item:
                        formatted.append(item)
            return formatted
        return []

    def format_projects(self, projects: Any) -> List[Union[Dict, str]]:
        """Format project entries into a consistent structure."""
        if not projects:
            return []
        if isinstance(projects, str):
            projects = projects.splitlines()
        formatted = []
        if isinstance(projects, (list, tuple, set)):
            for project in projects:
                if isinstance(project, dict):
                    entry = deepcopy(project)
                    entry["description"] = self.clean_text(entry.get("description") or entry.get("details") or "")
                    technologies = entry.get("technologies") or entry.get("tech_stack") or entry.get("skills") or []
                    entry["technologies"] = self.normalize_list(technologies)
                    formatted.append(entry)
                else:
                    project = self.clean_text(project)
                    if project:
                        formatted.append(project)
        return formatted

    def format_education(self, education: Any) -> List[Union[Dict, str]]:
        """Format education entries."""
        if not education:
            return []
        if isinstance(education, str):
            education = education.splitlines()
        formatted = []
        if isinstance(education, (list, tuple, set)):
            for item in education:
                if isinstance(item, dict):
                    formatted.append(deepcopy(item))
                else:
                    item = self.clean_text(item)
                    if item:
                        formatted.append(item)
        return formatted

    def format_certifications(self, certifications: Any) -> List[Union[Dict, str]]:
        """Format certification entries."""
        if not certifications:
            return []
        if isinstance(certifications, str):
            certifications = certifications.splitlines()
        formatted = []
        if isinstance(certifications, (list, tuple, set)):
            for item in certifications:
                if isinstance(item, dict):
                    formatted.append(deepcopy(item))
                else:
                    item = self.clean_text(item)
                    if item:
                        formatted.append(item)
        return formatted

    # ============================================================
    # PROTECTED INFORMATION EXTRACTION
    # ============================================================

    def extract_personal_info(self, resume: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract protected personal information from either:
        1. resume["personal_info"]
        2. top-level ResumeTailor output

        Values are copied, not generated.
        """
        if not isinstance(resume, dict):
            return {}

        source = resume.get("personal_info")
        if not isinstance(source, dict):
            source = resume

        personal_info = {}
        for field in self.protected_fields:
            if field in source:
                value = source.get(field)
                if value is not None:
                    if isinstance(value, (dict, list, tuple)):
                        personal_info[field] = deepcopy(value)
                    else:
                        personal_info[field] = str(value)
        return personal_info

    # ============================================================
    # PROTECTED INFORMATION VALIDATION
    # ============================================================

    def validate_protected_information(
        self,
        original: Dict[str, Any],
        generated: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Verify protected information remains unchanged.

        Returns:
            {
                "valid": True/False,
                "fields": {...}
            }
        """
        original_info = self.extract_personal_info(original)
        generated_info = self.extract_personal_info(generated)

        validation = {"valid": True, "fields": {}}
        for field in self.protected_fields:
            original_value = original_info.get(field, "")
            generated_value = generated_info.get(field, "")
            same = str(original_value) == str(generated_value)
            validation["fields"][field] = same
            if not same:
                validation["valid"] = False
        return validation

    # ============================================================
    # STRUCTURED RESUME
    # ============================================================

    def generate_resume(
        self,
        job_title: str = "",
        skills: Any = None,
        experience: int = 0,
        experience_details: Any = None,
        projects: Any = None,
        education: Any = None,
        certifications: Any = None,
        professional_summary: Optional[str] = None,
        personal_info: Optional[Dict[str, Any]] = None,
        source_resume: Optional[Dict[str, Any]] = None,
        missing_skills: Any = None,
        matched_skills: Any = None,
        **kwargs
    ) -> Dict[str, Any]:
        """
        Create structured resume data.

        Compatible with both:
        - ResumeGenerator input
        - ResumeTailor output

        IMPORTANT: missing_skills are never placed into the final skills.
        """
        # --------------------------------------------------------
        # SOURCE RESUME SUPPORT
        # --------------------------------------------------------
        if source_resume is not None:
            if not isinstance(source_resume, dict):
                raise ValueError("source_resume must be a dictionary.")

            source_personal_info = self.extract_personal_info(source_resume)
            if not personal_info:
                personal_info = source_personal_info

            if not job_title:
                job_title = source_resume.get("job_title") or ""

            if skills is None:
                skills = source_resume.get("skills") or []

            if experience_details is None:
                experience_details = source_resume.get("experience") or []

            if projects is None:
                projects = source_resume.get("projects") or []

            if education is None:
                education = source_resume.get("education") or []

            if certifications is None:
                certifications = source_resume.get("certifications") or []

            if professional_summary is None:
                professional_summary = source_resume.get("professional_summary") or source_resume.get("summary") or ""

            if missing_skills is None:
                missing_skills = source_resume.get("missing_skills") or []

            if matched_skills is None:
                matched_skills = source_resume.get("matched_skills") or []

        # --------------------------------------------------------
        # DIRECT RESUME TAILOR OUTPUT SUPPORT
        # --------------------------------------------------------
        if isinstance(personal_info, dict):
            personal_info = deepcopy(personal_info)
        else:
            personal_info = {}

        if not personal_info and kwargs:
            personal_info = self.extract_personal_info(kwargs)

        # --------------------------------------------------------
        # JOB TITLE
        # --------------------------------------------------------
        job_title = self.clean_text(job_title)

        # --------------------------------------------------------
        # SKILLS
        # --------------------------------------------------------
        normalized_skills = self.normalize_skills(skills or [])

        # --------------------------------------------------------
        # SUMMARY
        # --------------------------------------------------------
        if professional_summary:
            summary = self.clean_text(professional_summary)
        else:
            summary = self.generate_summary(
                job_title=job_title,
                skills=normalized_skills,
                experience=experience
            )

        # --------------------------------------------------------
        # FINAL STRUCTURE
        # --------------------------------------------------------
        result = {
            "personal_info": personal_info,
            "job_title": job_title,
            "professional_summary": summary,
            "skills": normalized_skills,
            "matched_skills": self.normalize_list(matched_skills),
            "missing_skills": self.normalize_list(missing_skills),
            "experience_years": experience,
            "experience": self.format_experience(experience_details),
            "projects": self.format_projects(projects),
            "education": self.format_education(education),
            "certifications": self.format_certifications(certifications),
        }

        # Optional supplied sections
        for section_name in ["achievements", "languages"]:
            if section_name in kwargs:
                result[section_name] = deepcopy(kwargs.get(section_name))

        return result

    # ============================================================
    # DOCUMENT STYLES
    # ============================================================

    def setup_styles(self, document: Document) -> None:
        """Set up document styles for a professional look."""
        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = self.default_font
        normal.font.size = Pt(9.5)

        if "Resume Section" not in styles:
            section_style = styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = self.default_font
            section_style.font.size = Pt(11)
            section_style.font.bold = True

        if "Resume Bullet" not in styles:
            bullet_style = styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
            bullet_style.font.name = self.default_font
            bullet_style.font.size = Pt(9.5)

    # ============================================================
    # PARAGRAPH BORDER
    # ============================================================

    @staticmethod
    def add_bottom_border(paragraph) -> None:
        """Add a thin bottom border to a paragraph."""
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = OxmlElement("w:pBdr")
            pPr.append(pBdr)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "808080")
        pBdr.append(bottom)

    # ============================================================
    # ADD HEADING
    # ============================================================

    def add_heading(self, document: Document, text: str) -> Optional[Paragraph]:
        """Add a section heading with a bottom border."""
        text = self.clean_text(text)
        if not text:
            return None
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(text.upper())
        run.bold = True
        run.font.name = self.default_font
        run.font.size = Pt(11)
        self.add_bottom_border(paragraph)
        return paragraph

    # ============================================================
    # ADD BULLET
    # ============================================================

    def add_bullet(self, document: Document, text: str) -> Optional[Paragraph]:
        """Add a bullet point paragraph."""
        text = self.clean_text(text)
        if not text:
            return None
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.12)
        paragraph.paragraph_format.space_after = Pt(2)

        bullet_run = paragraph.add_run("• ")
        bullet_run.font.name = self.default_font
        bullet_run.font.size = Pt(9.5)

        text_run = paragraph.add_run(text)
        text_run.font.name = self.default_font
        text_run.font.size = Pt(9.5)
        return paragraph

    # ============================================================
    # ADD PERSONAL INFORMATION
    # ============================================================

    def add_contact_information(self, document: Document, personal_info: Dict[str, Any]) -> None:
        """Render protected personal information (no AI rewriting)."""
        if not isinstance(personal_info, dict):
            return

        # Name
        name = personal_info.get("name", "")
        if name:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(str(name))
            run.bold = True
            run.font.name = self.default_font
            run.font.size = Pt(20)

        # Contact fields
        contact_values = []
        for key in ["email", "phone", "location", "linkedin", "github", "kaggle"]:
            value = personal_info.get(key)
            if value is None:
                continue
            value = str(value)
            if not value:
                continue
            display_value = self.clean_text(value)
            if display_value:
                contact_values.append(display_value)

        if not contact_values:
            return

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(6)
        for idx, value in enumerate(contact_values):
            run = paragraph.add_run(value)
            run.font.name = self.default_font
            run.font.size = Pt(8.5)
            if idx < len(contact_values) - 1:
                sep = paragraph.add_run("  |  ")
                sep.font.name = self.default_font
                sep.font.size = Pt(8.5)

    # ============================================================
    # ADD TARGET TITLE
    # ============================================================

    def add_target_title(self, document: Document, job_title: str) -> None:
        """Add the target job title."""
        job_title = self.clean_text(job_title)
        if not job_title:
            return
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(5)
        run = paragraph.add_run(job_title)
        run.bold = True
        run.font.name = self.default_font
        run.font.size = Pt(11)

    # ============================================================
    # ADD SUMMARY
    # ============================================================

    def add_summary(self, document: Document, summary: str) -> None:
        """Add the professional summary section."""
        summary = self.clean_text(summary)
        if not summary:
            return
        self.add_heading(document, "Professional Summary")
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(summary)
        run.font.name = self.default_font
        run.font.size = Pt(9.5)

    # ============================================================
    # ADD SKILLS
    # ============================================================

    def add_skills(self, document: Document, skills: Any) -> None:
        """Add skills section (categorized or plain list)."""
        if not skills:
            return

        self.add_heading(document, "Skills")

        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                skill_list = self.normalize_list(skill_list)
                skill_list = [s for s in skill_list if isinstance(s, str) and s.strip()]
                if not skill_list:
                    continue
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(2)
                category_run = paragraph.add_run(f"{self.clean_text(category)}: ")
                category_run.bold = True
                category_run.font.name = self.default_font
                category_run.font.size = Pt(9.5)
                skill_run = paragraph.add_run(", ".join(skill_list))
                skill_run.font.name = self.default_font
                skill_run.font.size = Pt(9.5)
        else:
            skill_list = self.normalize_list(skills)
            if skill_list:
                paragraph = document.add_paragraph()
                run = paragraph.add_run(", ".join(skill_list))
                run.font.name = self.default_font
                run.font.size = Pt(9.5)

    # ============================================================
    # ADD EXPERIENCE
    # ============================================================

    def add_experience(self, document: Document, experience: Any) -> None:
        """Add experience section (structured or simple)."""
        if not experience:
            return

        self.add_heading(document, "Experience")

        for item in experience:
            if isinstance(item, str):
                self.add_bullet(document, item)
                continue

            if not isinstance(item, dict):
                continue

            title = item.get("title") or item.get("role") or item.get("position") or ""
            company = item.get("company") or item.get("organization") or ""
            location = item.get("location") or ""
            dates = item.get("dates") or item.get("duration") or item.get("period") or ""

            header_parts = []
            if title:
                header_parts.append(self.clean_text(title))
            if company:
                header_parts.append(self.clean_text(company))
            if header_parts:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(1)
                run = paragraph.add_run(" | ".join(header_parts))
                run.bold = True
                run.font.name = self.default_font
                run.font.size = Pt(10)

            meta_parts = []
            if location:
                meta_parts.append(self.clean_text(location))
            if dates:
                meta_parts.append(self.clean_text(dates))
            if meta_parts:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(2)
                run = paragraph.add_run(" | ".join(meta_parts))
                run.italic = True
                run.font.name = self.default_font
                run.font.size = Pt(8.5)

            bullets = item.get("bullets") or item.get("description") or item.get("responsibilities") or []
            bullets = self.convert_to_bullets(bullets)
            for bullet in bullets:
                self.add_bullet(document, bullet)

    # ============================================================
    # ADD PROJECTS
    # ============================================================

    def add_projects(self, document: Document, projects: Any) -> None:
        """Add projects section (structured or simple)."""
        if not projects:
            return

        self.add_heading(document, "Projects")

        for project in projects:
            if isinstance(project, str):
                self.add_bullet(document, project)
                continue

            if not isinstance(project, dict):
                continue

            name = project.get("name") or project.get("title") or "Project"
            description = project.get("description") or project.get("details") or ""
            technologies = project.get("technologies") or project.get("tech_stack") or project.get("skills") or []

            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(1)
            name_run = paragraph.add_run(self.clean_text(name))
            name_run.bold = True
            name_run.font.name = self.default_font
            name_run.font.size = Pt(10)

            if description:
                self.add_bullet(document, description)

            technologies = self.normalize_list(technologies)
            if technologies:
                self.add_bullet(document, "Technologies: " + ", ".join(technologies))

    # ============================================================
    # ADD EDUCATION
    # ============================================================

    def add_education(self, document: Document, education: Any) -> None:
        """Add education section."""
        if not education:
            return

        self.add_heading(document, "Education")

        for item in education:
            if isinstance(item, str):
                self.add_bullet(document, item)
                continue

            if not isinstance(item, dict):
                continue

            degree = item.get("degree") or item.get("title") or item.get("qualification") or ""
            institution = item.get("institution") or item.get("university") or item.get("school") or ""
            dates = item.get("dates") or item.get("duration") or item.get("period") or ""

            parts = []
            if degree:
                parts.append(self.clean_text(degree))
            if institution:
                parts.append(self.clean_text(institution))
            if dates:
                parts.append(self.clean_text(dates))

            if parts:
                self.add_bullet(document, " | ".join(parts))

    # ============================================================
    # ADD CERTIFICATIONS
    # ============================================================

    def add_certifications(self, document: Document, certifications: Any) -> None:
        """Add certifications section."""
        if not certifications:
            return

        self.add_heading(document, "Certifications")

        for certificate in certifications:
            if isinstance(certificate, str):
                self.add_bullet(document, certificate)
                continue

            if not isinstance(certificate, dict):
                continue

            name = certificate.get("name") or certificate.get("title") or certificate.get("certificate") or ""
            issuer = certificate.get("issuer") or certificate.get("organization") or ""
            date = certificate.get("date") or certificate.get("year") or ""

            parts = []
            if name:
                parts.append(self.clean_text(name))
            if issuer:
                parts.append(self.clean_text(issuer))
            if date:
                parts.append(self.clean_text(date))

            if parts:
                self.add_bullet(document, " | ".join(parts))

    # ============================================================
    # OPTIONAL SECTION
    # ============================================================

    def add_optional_section(self, document: Document, title: str, items: Any) -> None:
        """Add an optional list-based section (e.g., Achievements, Languages)."""
        if not items:
            return
        self.add_heading(document, title)
        normalized = self.normalize_list(items)
        for item in normalized:
            if isinstance(item, dict):
                parts = []
                for key, value in item.items():
                    if value:
                        parts.append(f"{key}: {value}")
                text = " | ".join(parts)
            else:
                text = str(item)
            if text:
                self.add_bullet(document, text)

    # ============================================================
    # PAGE NUMBER
    # ============================================================

    @staticmethod
    def add_page_number(section) -> None:
        """Add a page number field to the footer."""
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run("Page ")
        run.font.name = "Arial"
        run.font.size = Pt(8)

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    # ============================================================
    # GENERATE DOCX
    # ============================================================

    def generate_docx(
        self,
        resume: Dict[str, Any],
        output_path: str = "output/Tailored_Resume.docx"
    ) -> str:
        """
        Generate an ATS-friendly DOCX resume from the supplied data.

        Supports both ResumeGenerator structure and ResumeTailor-compatible structure.

        Args:
            resume: Dictionary containing resume data.
            output_path: Path where the DOCX file will be saved.

        Returns:
            The absolute path to the generated file.
        """
        if not isinstance(resume, dict):
            raise ValueError("resume must be a dictionary.")

        # --------------------------------------------------------
        # DOCUMENT
        # --------------------------------------------------------
        document = Document()

        # --------------------------------------------------------
        # PAGE SETTINGS
        # --------------------------------------------------------
        section = document.sections[0]
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

        # --------------------------------------------------------
        # STYLES
        # --------------------------------------------------------
        self.setup_styles(document)

        # --------------------------------------------------------
        # PAGE NUMBER
        # --------------------------------------------------------
        self.add_page_number(section)

        # --------------------------------------------------------
        # PERSONAL INFORMATION
        # --------------------------------------------------------
        personal_info = self.extract_personal_info(resume)
        self.add_contact_information(document, personal_info)

        # --------------------------------------------------------
        # TARGET POSITION
        # --------------------------------------------------------
        job_title = resume.get("job_title") or resume.get("target_position") or "Professional"
        self.add_target_title(document, job_title)

        # --------------------------------------------------------
        # SUMMARY
        # --------------------------------------------------------
        self.add_summary(document, resume.get("professional_summary", ""))

        # --------------------------------------------------------
        # SKILLS
        # --------------------------------------------------------
        self.add_skills(document, resume.get("skills", {}))

        # --------------------------------------------------------
        # EXPERIENCE
        # --------------------------------------------------------
        self.add_experience(document, resume.get("experience", []))

        # --------------------------------------------------------
        # PROJECTS
        # --------------------------------------------------------
        self.add_projects(document, resume.get("projects", []))

        # --------------------------------------------------------
        # EDUCATION
        # --------------------------------------------------------
        self.add_education(document, resume.get("education", []))

        # --------------------------------------------------------
        # CERTIFICATIONS
        # --------------------------------------------------------
        self.add_certifications(document, resume.get("certifications", []))

        # --------------------------------------------------------
        # OPTIONAL SECTIONS
        # --------------------------------------------------------
        self.add_optional_section(document, "Achievements", resume.get("achievements", []))
        self.add_optional_section(document, "Languages", resume.get("languages", []))

        # --------------------------------------------------------
        # OUTPUT DIRECTORY
        # --------------------------------------------------------
        output_directory = os.path.dirname(output_path)
        if output_directory:
            os.makedirs(output_directory, exist_ok=True)

        # --------------------------------------------------------
        # SAVE
        # --------------------------------------------------------
        document.save(output_path)
        return output_path

    # ============================================================
    # PRINT RESUME
    # ============================================================

    def print_resume(self, resume: Dict[str, Any]) -> None:
        """Print a human-readable version of the resume to the console."""
        print()
        print("=" * 70)
        print("TAILORED RESUME")
        print("=" * 70)

        # PERSONAL INFORMATION
        personal_info = self.extract_personal_info(resume)
        if personal_info:
            print("\nPERSONAL INFORMATION")
            print("-" * 50)
            for key in self.protected_fields:
                value = personal_info.get(key)
                if value:
                    print(f"{key.capitalize()}: {value}")

        # TARGET POSITION
        print("\nTARGET POSITION")
        print("-" * 50)
        print(resume.get("job_title", "Professional"))

        # SUMMARY
        print("\nPROFESSIONAL SUMMARY")
        print("-" * 50)
        print(resume.get("professional_summary", ""))

        # SKILLS
        print("\nSKILLS")
        print("-" * 50)
        skills = resume.get("skills", {})
        if skills:
            if isinstance(skills, dict):
                for category, skill_list in skills.items():
                    values = self.normalize_list(skill_list)
                    print(f"{category}: " + ", ".join(str(x) for x in values))
            else:
                print(", ".join(self.normalize_list(skills)))
        else:
            print("No skills available.")

        # MATCHED SKILLS
        matched = resume.get("matched_skills", [])
        if matched:
            print("\nMATCHED JD SKILLS")
            print("-" * 50)
            print(", ".join(self.normalize_list(matched)))

        # MISSING SKILLS
        missing = resume.get("missing_skills", [])
        if missing:
            print("\nMISSING JD SKILLS")
            print("-" * 50)
            print(", ".join(self.normalize_list(missing)))
            print("\nNOTE: Missing skills are NOT added to the candidate resume.")

        # EXPERIENCE
        print("\nEXPERIENCE")
        print("-" * 50)
        experience = resume.get("experience", [])
        if experience:
            for item in experience:
                if isinstance(item, dict):
                    title = item.get("title") or item.get("role") or item.get("position") or ""
                    company = item.get("company") or item.get("organization") or ""
                    if title or company:
                        print(f"{title}" + (f" | {company}" if company else ""))
                    bullets = item.get("bullets") or item.get("description") or item.get("responsibilities") or []
                    for bullet in self.convert_to_bullets(bullets):
                        print("•", bullet)
                else:
                    print("•", item)
        else:
            print("No experience details available.")

        # PROJECTS
        print("\nPROJECTS")
        print("-" * 50)
        projects = resume.get("projects", [])
        if projects:
            for project in projects:
                if isinstance(project, dict):
                    name = project.get("name") or project.get("title") or "Project"
                    print("•", name)
                    description = project.get("description") or project.get("details") or ""
                    if description:
                        print("  ", description)
                    technologies = project.get("technologies") or project.get("tech_stack") or project.get("skills") or []
                    if technologies:
                        print("  Technologies:", ", ".join(self.normalize_list(technologies)))
                else:
                    print("•", project)
        else:
            print("No projects available.")

        # EDUCATION
        print("\nEDUCATION")
        print("-" * 50)
        education = resume.get("education", [])
        if education:
            for item in education:
                if isinstance(item, dict):
                    parts = []
                    for key in ["degree", "title", "qualification", "institution", "university", "school", "dates", "duration"]:
                        value = item.get(key, "")
                        if value:
                            parts.append(self.clean_text(value))
                    if parts:
                        print("•", " | ".join(parts))
                else:
                    print("•", item)
        else:
            print("No education information available.")

        # CERTIFICATIONS
        print("\nCERTIFICATIONS")
        print("-" * 50)
        certifications = resume.get("certifications", [])
        if certifications:
            for certificate in certifications:
                if isinstance(certificate, dict):
                    parts = []
                    for key in ["name", "title", "certificate", "issuer", "organization", "date", "year"]:
                        value = certificate.get(key, "")
                        if value:
                            parts.append(self.clean_text(value))
                    if parts:
                        print("•", " | ".join(parts))
                else:
                    print("•", certificate)
        else:
            print("No certifications available.")

        print()
        print("=" * 70)
        print("RESUME GENERATION COMPLETED")
        print("=" * 70)


# ================================================================
# TEST
# ================================================================

if __name__ == "__main__":
    print("=" * 70)
    print("PROFESSIONAL RESUME GENERATOR TEST")
    print("=" * 70)

    generator = ResumeGenerator()

    # ------------------------------------------------------------
    # PROTECTED INFORMATION
    # ------------------------------------------------------------
    personal_info = {
        "name": "Summaiya Bibi",
        "email": "summaiyabibi4545@gmail.com",
        "phone": "03466577540",
        "location": "Pakistan",
        "linkedin": "https://linkedin.com/in/example",
        "github": "https://github.com/summaiyazafar",
        "kaggle": "https://kaggle.com/example",
    }

    # ------------------------------------------------------------
    # EXISTING SKILLS
    # ------------------------------------------------------------
    skills = [
        "Python",
        "SQL",
        "Pandas",
        "NumPy",
        "Machine Learning",
        "Power BI",
        "Scikit-learn",
    ]

    # ------------------------------------------------------------
    # EXPERIENCE
    # ------------------------------------------------------------
    experience_details = [
        {
            "title": "Machine Learning / Data Science Experience",
            "company": "Existing Experience",
            "dates": "Existing Dates",
            "bullets": [
                "Developed machine learning models for predictive analytics.",
                "Analyzed datasets using Python, Pandas and NumPy.",
                "Created data visualizations and dashboards using Power BI.",
            ],
        }
    ]

    # ------------------------------------------------------------
    # PROJECTS
    # ------------------------------------------------------------
    projects = [
        {
            "name": "Machine Learning Prediction System",
            "description": "Developed a machine learning prediction system using Python and Scikit-learn.",
            "technologies": ["Python", "Scikit-learn", "Pandas", "NumPy"],
        },
        {
            "name": "Data Analytics Dashboard",
            "description": "Created an analytics dashboard for data visualization and reporting.",
            "technologies": ["Power BI", "SQL", "Excel"],
        },
    ]

    # ------------------------------------------------------------
    # EDUCATION
    # ------------------------------------------------------------
    education = [
        {
            "degree": "BS Computer Science",
            "institution": "Virtual University of Pakistan",
        }
    ]

    # ------------------------------------------------------------
    # CERTIFICATIONS
    # ------------------------------------------------------------
    certifications = [
        "Artificial Intelligence using Python",
        "Data Analytics & Business Intelligence",
    ]

    # ------------------------------------------------------------
    # GENERATE
    # ------------------------------------------------------------
    resume = generator.generate_resume(
        job_title="Machine Learning Engineer",
        skills=skills,
        experience_details=experience_details,
        projects=projects,
        education=education,
        certifications=certifications,
        professional_summary=(
            "Machine learning and data science professional "
            "with practical experience in Python, SQL, "
            "Pandas, NumPy, Power BI and Scikit-learn."
        ),
        personal_info=personal_info,
        matched_skills=["Python", "SQL", "Machine Learning"],
        missing_skills=["Azure", "Kubernetes"],
    )

    # ------------------------------------------------------------
    # PRINT
    # ------------------------------------------------------------
    generator.print_resume(resume)

    # ------------------------------------------------------------
    # VALIDATE PROTECTED INFORMATION
    # ------------------------------------------------------------
    validation = generator.validate_protected_information(personal_info, resume)
    print()
    print("PROTECTED INFORMATION VALIDATION")
    print("-" * 50)
    print(validation)

    # ------------------------------------------------------------
    # GENERATE DOCX
    # ------------------------------------------------------------
    output_file = generator.generate_docx(resume, "output/Tailored_Resume.docx")
    print()
    print("DOCX FILE CREATED:")
    print(output_file)