"""
AI Resume Tailoring System – Resume Parser
===========================================

Extracts protected and editable information from PDF, DOCX, or TXT resumes.
"""

import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Union

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None


class ResumeParser:
    """
    Resume parser that extracts protected and editable information.
    Protected fields (name, contact, social URLs, education, certifications)
    are preserved exactly as extracted.
    """

    def __init__(self) -> None:
        self.supported_extensions: Set[str] = {".pdf", ".docx", ".txt"}
        self.protected_fields: Set[str] = {
            "name", "phone", "email", "linkedin", "github",
            "kaggle", "location", "education", "certifications"
        }
        self.editable_fields: Set[str] = {
            "summary", "skills", "experience", "projects"
        }

    # ==========================================================
    # FILE PARSING
    # ==========================================================

    def parse_pdf(self, file_path: Union[str, Path]) -> str:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Resume not found: {file_path}")
        if PdfReader is None:
            raise ImportError("pip install pypdf")

        try:
            reader = PdfReader(file_path)
        except Exception as exc:
            raise ValueError(f"Unable to read PDF: {exc}")

        text_parts = []
        for page in reader.pages:
            try:
                page_text = page.extract_text()
            except Exception:
                page_text = ""
            if page_text:
                text_parts.append(page_text)

        text = "\n".join(text_parts)
        if not text.strip():
            raise ValueError("No readable text found in PDF.")
        return text

    def parse_docx(self, file_path: Union[str, Path]) -> str:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Resume not found: {file_path}")
        if Document is None:
            raise ImportError("pip install python-docx")

        try:
            document = Document(file_path)
        except Exception as exc:
            raise ValueError(f"Unable to read DOCX: {exc}")

        content = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                content.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    content.append(" | ".join(cells))
        return "\n".join(content)

    def parse_txt(self, file_path: Union[str, Path]) -> str:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Resume not found: {file_path}")

        encodings = ["utf-8", "utf-8-sig", "cp1252", "latin-1"]
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding, errors="ignore") as f:
                    return f.read()
            except Exception:
                continue
        raise ValueError("Unable to read TXT file.")

    def parse(self, file_path: Union[str, Path]) -> str:
        if not file_path:
            raise ValueError("Resume file path is required.")
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Resume not found: {file_path}")

        ext = file_path.suffix.lower()
        if ext not in self.supported_extensions:
            raise ValueError("Unsupported format. Use PDF, DOCX, or TXT.")

        if ext == ".pdf":
            text = self.parse_pdf(file_path)
        elif ext == ".docx":
            text = self.parse_docx(file_path)
        else:
            text = self.parse_txt(file_path)

        return self.clean_text(text)

    # ==========================================================
    # TEXT CLEANING
    # ==========================================================

    @staticmethod
    def clean_text(text: str) -> str:
        if not text:
            return ""
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = text.replace("\xa0", " ").replace("\u200b", "").replace("\ufeff", "")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"[ \t]+\n", "\n", text)
        text = re.sub(r"\n[ \t]+", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def normalize_line(line: str) -> str:
        if not line:
            return ""
        line = str(line).strip().lower()
        line = line.replace("**", "").replace("__", "")
        line = re.sub(r"[:|]+$", "", line)
        return re.sub(r"\s+", " ", line).strip()

    # ==========================================================
    # PROTECTED EXTRACTION
    # ==========================================================

    @staticmethod
    def extract_email(text: str) -> str:
        if not text:
            return ""
        match = re.search(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b", text)
        return match.group(0).strip() if match else ""

    @staticmethod
    def extract_phone(text: str) -> str:
        if not text:
            return ""
        patterns = [
            r"(?:\+92|0092)[\s\-]?3\d{2}[\s\-]?\d{7}",
            r"\b03\d{2}[\s\-]\d{7}\b",
            r"\b03\d{9}\b",
            r"\b\d{11}\b"
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0).strip()
        return ""

    @staticmethod
    def normalize_url(url: str) -> str:
        if not url:
            return ""
        url = str(url).strip().strip(" \t\r\n.,;:()[]{}<>")
        if not url:
            return ""
        if not url.lower().startswith(("http://", "https://")):
            url = "https://" + url
        return url

    def extract_linkedin(self, text: str) -> str:
        if not text:
            return ""
        for pattern in [
            r"https?://(?:www\.)?linkedin\.com/in/[A-Za-z0-9._%\-]+",
            r"(?:www\.)?linkedin\.com/in/[A-Za-z0-9._%\-]+"
        ]:
            match = re.search(pattern, text, flags=re.IGNORECASE)
            if match:
                return self.normalize_url(match.group(0).rstrip("/.,;:)]}"))
        return ""

    def extract_github(self, text: str) -> str:
        if not text:
            return ""
        for pattern in [
            r"https?://(?:www\.)?github\.com/[A-Za-z0-9._\-]+",
            r"(?:www\.)?github\.com/[A-Za-z0-9._\-]+"
        ]:
            match = re.search(pattern, text, flags=re.IGNORECASE)
            if match:
                return self.normalize_url(match.group(0).rstrip("/.,;:)]}"))
        return ""

    def extract_kaggle(self, text: str) -> str:
        if not text:
            return ""
        for pattern in [
            r"https?://(?:www\.)?kaggle\.com/[A-Za-z0-9._\-]+",
            r"(?:www\.)?kaggle\.com/[A-Za-z0-9._\-]+"
        ]:
            match = re.search(pattern, text, flags=re.IGNORECASE)
            if match:
                return self.normalize_url(match.group(0).rstrip("/.,;:)]}"))
        return ""

    def extract_location(self, text: str) -> str:
        if not text:
            return ""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not lines:
            return ""

        for line in lines[:20]:
            lower = line.lower()
            for label in ["location", "address", "based in", "city"]:
                match = re.search(r"^" + re.escape(label) + r"\s*[:\-]\s*(.+)$", lower, flags=re.IGNORECASE)
                if match:
                    value = line[match.start(1):].strip()
                    if value:
                        return value

        known = ["islamabad", "rawalpindi", "lahore", "karachi", "peshawar",
                 "multan", "faisalabad", "taxila", "attock", "hazro",
                 "wah cantt", "quetta", "sialkot", "gujranwala",
                 "punjab", "sindh", "khyber pakhtunkhwa", "balochistan"]
        for line in lines[:20]:
            lower = line.lower()
            for loc in known:
                if loc in lower and len(line) <= 100:
                    return line
        return ""

    def extract_name(self, text: str) -> str:
        if not text:
            return ""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not lines:
            return ""

        ignored_exact = {"resume", "curriculum vitae", "cv", "curriculum-vitae",
                         "profile", "contact", "personal information", "personal details"}
        ignored_contains = ["email", "phone", "mobile", "linkedin", "github",
                            "kaggle", "portfolio", "http://", "https://", "@", "www."]
        title_words = {"data analyst", "data scientist", "machine learning engineer",
                       "ai engineer", "ml engineer", "software engineer", "python developer",
                       "web developer", "full stack developer", "backend developer",
                       "frontend developer", "artificial intelligence", "machine learning",
                       "data science", "developer", "engineer", "analyst", "scientist"}
        section_words = {"skills", "experience", "education", "certifications",
                         "projects", "summary", "objective", "professional",
                         "technical", "contact", "achievements", "languages"}

        for line in lines[:15]:
            lower = line.lower().strip()
            if lower in ignored_exact:
                continue
            if any(term in lower for term in ignored_contains):
                continue
            if lower in title_words:
                continue
            if "|" in line or " - " in line:
                parts = line.split(" - ") if " - " in line else line.split("|")
                if any(part.strip().lower() in title_words for part in parts):
                    continue
            if len(line) > 60:
                continue
            words = line.split()
            if not (2 <= len(words) <= 5):
                continue
            if any(word.lower().strip(":") in section_words for word in words):
                continue
            valid = sum(1 for w in words if re.search(r'[A-Za-zÀ-ÿ]', w))
            if valid >= 2:
                job_titles = {"analyst", "engineer", "developer", "scientist",
                              "manager", "intern", "consultant"}
                if any(w.lower() in job_titles for w in words):
                    continue
                return line
        return ""

    # ==========================================================
    # SECTION EXTRACTION
    # ==========================================================

    def extract_section(self, text: str, section_names: List[str]) -> str:
        if not text:
            return ""

        lines = text.splitlines()
        target_names = {self.normalize_line(name) for name in section_names}

        start_index = None
        for i, line in enumerate(lines):
            if self.normalize_line(line) in target_names:
                start_index = i + 1
                break

        if start_index is None:
            return ""

        common_sections = {
            "professional summary", "summary", "profile", "professional profile",
            "objective", "career objective", "skills", "technical skills",
            "core skills", "key skills", "professional skills",
            "technical competencies", "areas of expertise", "technical expertise",
            "experience", "work experience", "professional experience",
            "employment history", "work history", "career history",
            "projects", "personal projects", "academic projects", "key projects",
            "project experience", "education", "academic background",
            "academic qualifications", "educational background",
            "certifications", "certificates", "professional certifications",
            "licenses & certifications", "achievements", "awards",
            "languages", "references"
        }

        result = []
        for line in lines[start_index:]:
            normalized = self.normalize_line(line)
            if normalized in common_sections and result:
                break
            if line.strip():
                result.append(line.strip())

        return "\n".join(result).strip()

    def extract_summary_section(self, text: str) -> str:
        return self.extract_section(text, [
            "professional summary", "summary", "profile",
            "professional profile", "career objective", "objective"
        ])

    def extract_skills_section(self, text: str) -> str:
        return self.extract_section(text, [
            "skills", "technical skills", "core skills", "key skills",
            "professional skills", "technical competencies",
            "areas of expertise", "technical expertise"
        ])

    def extract_experience_section(self, text: str) -> str:
        return self.extract_section(text, [
            "experience", "work experience", "professional experience",
            "employment history", "work history", "career history"
        ])

    def extract_projects_section(self, text: str) -> str:
        return self.extract_section(text, [
            "projects", "personal projects", "academic projects",
            "key projects", "project experience"
        ])

    def extract_education_section(self, text: str) -> str:
        return self.extract_section(text, [
            "education", "academic background", "academic qualifications",
            "educational background"
        ])

    def extract_certifications_section(self, text: str) -> str:
        return self.extract_section(text, [
            "certifications", "certificates", "professional certifications",
            "licenses & certifications", "license certifications",
            "certification"
        ])

    # ==========================================================
    # MAIN EXTRACTION
    # ==========================================================

    def extract_protected_information(self, text: str) -> Dict[str, str]:
        return {
            "name": self.extract_name(text),
            "phone": self.extract_phone(text),
            "email": self.extract_email(text),
            "linkedin": self.extract_linkedin(text),
            "github": self.extract_github(text),
            "kaggle": self.extract_kaggle(text),
            "location": self.extract_location(text),
            "education": self.extract_education_section(text),
            "certifications": self.extract_certifications_section(text),
        }

    def extract_information(self, text: str) -> Dict[str, Any]:
        text = self.clean_text(text)
        protected = self.extract_protected_information(text)
        editable = {
            "summary": self.extract_summary_section(text),
            "skills": self.extract_skills_section(text),
            "experience": self.extract_experience_section(text),
            "projects": self.extract_projects_section(text),
        }

        return {
            "text": text,
            "protected": protected,
            "editable": editable,
            "name": protected["name"],
            "phone": protected["phone"],
            "email": protected["email"],
            "linkedin": protected["linkedin"],
            "github": protected["github"],
            "kaggle": protected["kaggle"],
            "location": protected["location"],
            "education": protected["education"],
            "certifications": protected["certifications"],
            "summary": editable["summary"],
            "skills": editable["skills"],
            "experience": editable["experience"],
            "projects": editable["projects"],
        }

    def validate_protected_information(self, original: Dict, generated: Dict) -> Dict:
        errors = []
        for field in self.protected_fields:
            if str(original.get(field, "")).strip() != str(generated.get(field, "")).strip():
                errors.append(f"Protected field changed: {field}")
        return {"valid": len(errors) == 0, "errors": errors}

    def parse_resume(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        text = self.parse(file_path)
        return self.extract_information(text)


# ==============================================================
# TEST
# ==============================================================

if __name__ == "__main__":
    parser = ResumeParser()
    test_text = """
    Summaiya Bibi
    Email: summaiya@example.com
    Phone: 03001234567
    Location: Islamabad, Pakistan
    """
    result = parser.extract_information(test_text)
    print("\n" + "=" * 60)
    print("RESUME PARSER TEST")
    print("=" * 60)
    for key, value in result["protected"].items():
        print(f"{key}: {value}")
    print("=" * 60)